"""
CONTAINER INSPECTION REPORT WEB APPLICATION
Flask Backend Server
"""

from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, jsonify, send_from_directory
from werkzeug.utils import secure_filename
import os
import json
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from document_extractor import process_uploaded_documents, analyze_damage_image, extract_iauditor_images, narrate_iauditor_sections, extract_iauditor_file_summary
from email_service import GraphEmailService, EmailServiceError

app = Flask(__name__)
app.secret_key = 'your-secret-key-change-this-in-production'

# Configuration
UPLOAD_FOLDER = 'uploads'
REPORTS_FOLDER = 'reports'
DATA_FOLDER = 'data'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'jpg', 'jpeg', 'png'}
IMAGE_EXTENSIONS = {'jpg', 'jpeg', 'png'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['REPORTS_FOLDER'] = REPORTS_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORTS_FOLDER, exist_ok=True)
os.makedirs(DATA_FOLDER, exist_ok=True)

USERS = {
    'admin': 'admin123',
    'surveyor': 'survey123',
    'demo': 'demo123'
}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def allowed_image(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in IMAGE_EXTENSIONS


def load_case_data():
    cases_file = os.path.join(DATA_FOLDER, 'cases.json')
    if os.path.exists(cases_file):
        with open(cases_file, 'r') as f:
            return json.load(f)
    return []


def save_case_data(cases):
    cases_file = os.path.join(DATA_FOLDER, 'cases.json')
    with open(cases_file, 'w') as f:
        json.dump(cases, f, indent=2)


def format_date(date_str):
    """Normalize date to 'DD MONTH YYYY' format e.g. '08 APRIL 2024'."""
    if not date_str or date_str == 'N/A' or date_str == '#':
        return date_str or '#DATE'
    # Already in correct format like "08 APRIL 2024"
    from datetime import datetime as dt
    for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y', '%d-%m-%Y', '%Y/%m/%d'):
        try:
            parsed = dt.strptime(date_str.strip(), fmt)
            return parsed.strftime('%d %B %Y').upper()
        except ValueError:
            continue
    # Try partial formats like "4 Apr 2024" or "4Apr202400:00+04"
    import re
    match = re.match(r'(\d{1,2})\s*([A-Za-z]+)\s*(\d{4})', date_str)
    if match:
        day, month, year = match.groups()
        try:
            parsed = dt.strptime(f"{day} {month} {year}", '%d %b %Y')
            return parsed.strftime('%d %B %Y').upper()
        except ValueError:
            try:
                parsed = dt.strptime(f"{day} {month} {year}", '%d %B %Y')
                return parsed.strftime('%d %B %Y').upper()
            except ValueError:
                pass
    # Return as-is but uppercased if it looks like a date
    return date_str.upper() if any(m in date_str.upper() for m in ['JAN','FEB','MAR','APR','MAY','JUN','JUL','AUG','SEP','OCT','NOV','DEC']) else date_str


def find_case_by_id(cases, case_id):
    """Find a case by its unique ID (not case_reference, which can duplicate)."""
    return next((c for c in cases if c.get('id') == case_id), None)


def find_case_by_ref(cases, case_ref):
    """Find most recent case by reference (fallback for old data)."""
    # Search in reverse to get the most recent one
    for c in reversed(cases):
        if c.get('case_reference') == case_ref:
            return c
    return None


# ================================================================
# REPORT GENERATOR - ISA Certificate of Survey Template
# ================================================================

def _add_hyperlink(paragraph, url, text, size=11, color=RGBColor(0x05, 0x63, 0xC1)):
    """Append a clickable hyperlink run to a python-docx paragraph."""
    from docx.oxml.shared import OxmlElement, qn
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyper = OxmlElement('w:hyperlink')
    hyper.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rPr.append(sz)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rPr.append(rFonts)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    hyper.append(new_run)
    paragraph._p.append(hyper)
    return hyper


def generate_full_report(case_data):
    """Generate report following ISA Middle East / Quest Marine Certificate of Survey template."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Heading 1: Calibri Bold 14, Dark Blue
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)  # dark blue
    # Heading 2: Calibri Bold 13, Dark Blue
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)  # dark blue

    # Pre-format all date fields to "DD MONTH YYYY" format
    date_fields = [
        'report_date', 'survey_date', 'bl_issue_date', 'arrival_date',
        'delivery_date', 'gate_out_date', 'container_return_date',
        'vessel_loading_date', 'notice_of_loss_date', 'invoice_date',
        'collection_date', 'instruction_date', 'development_date',
        'transhipment_arrival_date', 'transhipment_reload_date',
    ]
    for df in date_fields:
        val = case_data.get(df, '')
        if val and val != 'N/A':
            case_data[df] = format_date(val)

    # Normalise weight fields — strip any trailing "KGS"/"KG" so we don't print "KGS KGS"
    import re as _re
    for _wf in ('gross_weight', 'net_weight'):
        _v = case_data.get(_wf, '')
        if isinstance(_v, str) and _v:
            case_data[_wf] = _re.sub(r'\s*kgs?\.?\s*$', '', _v, flags=_re.IGNORECASE).strip()

    case_ref = case_data.get('case_reference', 'N/A')
    report_date = case_data.get('report_date', datetime.now().strftime('%d %B %Y').upper())

    # ============================================================
    # PAGE 1: COVER PAGE - CERTIFICATE OF SURVEY
    # ============================================================

    # Persistent header on pages 2+ ("Our Ref" / "Your Ref"), suppressed on page 1
    principal_ref = case_data.get('principal_reference', 'N/A')
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    # Blank first-page header
    fp_hdr = section.first_page_header
    if fp_hdr.paragraphs:
        fp_hdr.paragraphs[0].text = ''
    hdr = section.header
    hdr_p1 = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hdr_p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hdr_p1.text = ''
    for txt in [("Our Ref : ", False), (str(case_ref), False)]:
        r = hdr_p1.add_run(txt[0])
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.bold = txt[1]
    hdr_p2 = hdr.add_paragraph()
    hdr_p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for txt in [("Your Ref : ", False), (str(principal_ref), False)]:
        r = hdr_p2.add_run(txt[0])
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.bold = txt[1]

    # Cover-page top block: logo (left) and address (right) at the same level,
    # using a borderless 2-column table.
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             'static', 'images', 'quest-marine-logo.png')
    top_table = doc.add_table(rows=1, cols=2)
    top_table.autofit = False
    top_table.columns[0].width = Inches(3.2)
    top_table.columns[1].width = Inches(3.3)
    left_cell, right_cell = top_table.rows[0].cells
    left_cell.width = Inches(3.2)
    right_cell.width = Inches(3.3)

    # Left: logo
    left_p = left_cell.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(logo_path):
        try:
            left_p.add_run().add_picture(logo_path, width=Inches(2.2))
        except Exception as e:
            print(f"Failed to embed logo: {e}")

    # Right: address block, top-aligned
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for line, size in [
        ("ISA Middle East", 10),
        ("c/o Quest Marine LLC", 10),
        ("P.O. Box : 123276", 10),
        ("Dubai, U.A.E", 10),
        ("", 10),
        ("Tel:        +971 (0) 4 4425206", 9),
        ("Mobile:   +971 (0) 56 1707676", 9),
        ("E-mail:   middle-east@isa-surveys.com", 9),
        ("Website: www.isa-surveys.com", 9),
    ]:
        if line == "" and right_p.text == "":
            run = right_p.add_run("\n")
        elif right_p.text == "" and not right_p.runs:
            run = right_p.add_run(line)
            run.font.size = Pt(size)
        else:
            run = right_p.add_run("\n" + line)
            run.font.size = Pt(size)

    doc.add_paragraph()

    # CERTIFICATE OF SURVEY title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("CERTIFICATE OF SURVEY")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()

    # Reference line - Calibri 13, black, single line via 3-col borderless table
    ref_font_size = Pt(13)
    ref_black = RGBColor(0x00, 0x00, 0x00)

    ref_table = doc.add_table(rows=1, cols=3)
    ref_table.autofit = True
    ref_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    ref_cells = ref_table.rows[0].cells

    def _fill_ref_cell(cell, label, value):
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        run_lbl = cell.paragraphs[0].add_run(label)
        run_lbl.bold = True
        run_lbl.font.name = 'Calibri'
        run_lbl.font.size = ref_font_size
        run_lbl.font.color.rgb = ref_black
        run_val = cell.paragraphs[0].add_run(value)
        run_val.font.name = 'Calibri'
        run_val.font.size = ref_font_size
        run_val.font.color.rgb = ref_black

    _fill_ref_cell(ref_cells[0], "OUR REF: ", f"{case_ref}")
    _fill_ref_cell(ref_cells[1], "YOUR REF: ", case_data.get('principal_reference', '#PRINCIPAL REF'))
    _fill_ref_cell(ref_cells[2], "DATE: ", f"{report_date}")

    # Strip table borders so it looks like a plain line
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl_pr = ref_table._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        ref_table._element.insert(0, tbl_pr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        borders.append(b)
    tbl_pr.append(borders)

    doc.add_paragraph()
    doc.add_paragraph()

    # Transit routing - multi-leg vessel info
    transit_routing = case_data.get('transit_routing', '')
    if transit_routing and transit_routing != 'N/A':
        for leg in transit_routing.split('\n'):
            leg = leg.strip()
            if leg:
                p = doc.add_paragraph()
                run = p.add_run(leg.upper())
                run.bold = True
    else:
        vessel_line = doc.add_paragraph()
        vessel_line.add_run(
            f"M/V \"{case_data.get('vessel_name', '#VESSEL NAME')}\" "
            f"VOY. {case_data.get('voyage_number', '#VOYAGE NO.')} "
            f"- TRANSIT FROM {case_data.get('origin_country', '#COUNTRY')} "
            f"TO {case_data.get('destination_country', '#COUNTRY')}"
        ).bold = True

    # BL and consignment summary
    doc.add_paragraph()
    bl_line = doc.add_paragraph()
    bl_line.add_run(
        f"{case_data.get('carrier_name', '#CARRIER').upper()} B/L NO. {case_data.get('bl_number', '#B/L NO.')}, "
        f"ISSUED ON {case_data.get('bl_issue_date', '#DATE').upper()}"
    ).bold = True

    consignment = doc.add_paragraph()
    num_cntrs = case_data.get('number_of_containers', '#NO.')
    cntr_tp = case_data.get('container_type', '#TYPE')
    cntr_tp_fmt = cntr_tp.replace("'", "FT").replace("ft", "FT").upper() if cntr_tp and cntr_tp != 'N/A' else cntr_tp
    consignment.add_run(
        f"CONSIGNMENT {num_cntrs} X {cntr_tp_fmt} CONTAINERS STC "
        f"{case_data.get('gross_weight', '#WEIGHT')} KGS OF "
        f"{case_data.get('goods_description', '#GOODS').upper()}"
    ).bold = True

    cntr = doc.add_paragraph()
    cntr.add_run(
        f"CONTAINER NOS. {case_data.get('container_number', '#CONTAINER NOS')}"
    ).bold = True

    shipper_line = doc.add_paragraph()
    shipper_line.add_run(
        f"SHIPPER: {case_data.get('shipper', '#SHIPPER').upper()}"
    ).bold = True

    consignee_line = doc.add_paragraph()
    consignee_line.add_run(
        f"CONSIGNEE: {case_data.get('consignee', '#CONSIGNEE').upper()}"
    ).bold = True

    doc.add_paragraph()

    # Instruction paragraph
    p = doc.add_paragraph()
    p.add_run(
        f"We hereby Certify that in accordance with the instructions of "
        f"{case_data.get('principal_name', '#PRINCIPAL\'S NAME & ADDRESS')}, "
        f"received on {case_data.get('instruction_date', '#DATE')}, we, the undersigned Marine Surveyors, "
        f"did immediately communicate with parties concerned and undertook a survey on the following:"
    )

    doc.add_paragraph()

    # ============================================================
    # PARTICULARS OF SURVEY table — start on page 2
    # ============================================================
    doc.add_page_break()
    survey_heading = doc.add_paragraph()
    run = survey_heading.add_run("PARTICULARS OF SURVEY")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

    survey_table = doc.add_table(rows=4, cols=2)
    survey_table.style = 'Table Grid'

    # Surveyor in Charge = attending_surveyor from iAuditor
    surveyor_name = case_data.get('attending_surveyor', '') or case_data.get('surveyor_in_charge', '#')
    # Place & Date = survey_location + survey_date from iAuditor
    survey_loc = case_data.get('survey_location', '')
    survey_dt = case_data.get('survey_date', '')
    if survey_loc and survey_loc != 'N/A' and survey_dt and survey_dt != 'N/A':
        survey_place_date = f"{survey_loc}, {survey_dt}"
    elif survey_loc and survey_loc != 'N/A':
        survey_place_date = survey_loc
    elif survey_dt and survey_dt != 'N/A':
        survey_place_date = survey_dt
    else:
        survey_place_date = '#'

    survey_fields = [
        ('Surveyor in Charge:', surveyor_name),
        ('Attending Surveyor:', surveyor_name),
        ('Place & Date of Survey(s):', survey_place_date),
        ('Other Parties in Attendance:', case_data.get('other_parties', '#')),
    ]

    for i, (label, value) in enumerate(survey_fields):
        survey_table.rows[i].cells[0].text = label
        for run in survey_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True
        survey_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ============================================================
    # PARTICULARS OF TRANSIT table
    # ============================================================
    transit_heading = doc.add_paragraph()
    run = transit_heading.add_run("PARTICULARS OF TRANSIT")
    run.bold = True
    run.font.size = Pt(11)

    transit_table = doc.add_table(rows=7, cols=2)
    transit_table.style = 'Table Grid'

    # Container number + type formatted as "3 X 40FT CONTAINERS"
    cntr_num = case_data.get('container_number', '#')
    cntr_type = case_data.get('container_type', '')
    num_containers = case_data.get('number_of_containers', '1')
    # Normalize type: "40' HC" -> "40FT HC"
    if cntr_type and cntr_type != 'N/A':
        cntr_type_fmt = cntr_type.replace("'", "FT").replace("ft", "FT").upper()
        container_info = f"{num_containers} X {cntr_type_fmt} CONTAINERS, NOS. {cntr_num}"
    else:
        container_info = cntr_num

    bl_info = f"{case_data.get('carrier_name', '#CARRIER')} Bill of Lading No. {case_data.get('bl_number', '#B/L NO.')} issued at {case_data.get('bl_issue_place', '#PLACE')} on {case_data.get('bl_issue_date', '#DATE')}"
    shipment_from_to = f"{case_data.get('origin_port', '#LOAD PORT')}, {case_data.get('origin_country', '#COUNTRY')} to {case_data.get('discharge_port', '#DISPORT')}, {case_data.get('destination_country', '#COUNTRY')}"

    # Delivery place/date from iAuditor "Date of delivery to the Consignee"
    delivery_date = case_data.get('delivery_date', '#DATE')
    delivery_loc = case_data.get('delivery_location', '')
    if not delivery_loc or delivery_loc == 'N/A':
        delivery_loc = case_data.get('survey_location', '')
    if not delivery_loc or delivery_loc == 'N/A':
        delivery_loc = ''
    if delivery_loc and delivery_date and delivery_date != '#DATE':
        delivery_info = f"{delivery_loc} on {delivery_date}"
    elif delivery_date and delivery_date != '#DATE':
        delivery_info = delivery_date
    elif delivery_loc:
        delivery_info = delivery_loc
    else:
        delivery_info = '#PLACE on #DATE'

    transit_fields = [
        ('Container Number(s) and type:', container_info),
        ('Bill of Lading No. and Date:', bl_info),
        ('Shipment - From / To:', shipment_from_to),
        ('Vessel Arrival Date:', case_data.get('arrival_date', '#')),
        ('Place and Date of Final Delivery:', delivery_info),
        ('Nature of Receipt Given on Delivery:', case_data.get('delivery_receipt', 'Delivery Receipt / EIR claused "#REMARKS"')),
        ('', ''),
    ]

    for i, (label, value) in enumerate(transit_fields):
        if label:
            transit_table.rows[i].cells[0].text = label
            for run in transit_table.rows[i].cells[0].paragraphs[0].runs:
                run.bold = True
            transit_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ============================================================
    # PARTICULARS OF CARGO table
    # ============================================================
    cargo_heading = doc.add_paragraph()
    run = cargo_heading.add_run("PARTICULARS OF CARGO")
    run.bold = True
    run.font.size = Pt(11)

    cargo_table = doc.add_table(rows=3, cols=2)
    cargo_table.style = 'Table Grid'

    cargo_fields = [
        ('Shippers:', case_data.get('shipper', '#NAME & ADDRESS')),
        ('Consignees:', case_data.get('consignee', '#NAME & ADDRESS')),
        ('Quantity and Description of Goods:', f"{case_data.get('goods_description', '#GOODS PER BILL OF LADING')}\nGW: {case_data.get('gross_weight', '#WEIGHT')} KGS  NW: {case_data.get('net_weight', '#WEIGHT')} KGS"),
    ]

    for i, (label, value) in enumerate(cargo_fields):
        cargo_table.rows[i].cells[0].text = label
        for run in cargo_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True
        cargo_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ============================================================
    # RESERVES AGAINST CARRIERS
    # ============================================================
    reserves_heading = doc.add_paragraph()
    run = reserves_heading.add_run("RESERVES AGAINST CARRIERS")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run(
        f"Notice of Loss was issued by {case_data.get('consignee', '#COMPANY')} to the Carrier, "
        f"{case_data.get('carrier_name', '#CARRIER')} on {case_data.get('notice_of_loss_date', '#DATE')}, "
        f"a copy of which is attached to this report."
    ).bold = True

    doc.add_paragraph()

    # ============================================================
    # 1. BACKGROUND
    # ============================================================
    doc.add_heading('1.    BACKGROUND', level=1)

    doc.add_heading('1.1    Circumstances Leading to Claim', level=2)

    p = doc.add_paragraph()
    p.add_run(
        f"From documentation and information made available, we understand that the subject consignment, "
        f"comprising {case_data.get('number_of_packages', '#PKGS')} of {case_data.get('goods_description', '#GOODS')}, "
        f"was sold by the Shipper, {case_data.get('shipper', '#NAME OF SHIPPER, #COUNTRY')} "
        f"to the Consignee, {case_data.get('consignee', '#NAME OF CONSIGNEE, #COUNTRY')} "
        f"on {case_data.get('incoterms', '#INCOTERMS')}."
    )

    p = doc.add_paragraph()
    p.add_run(
        f"According to information secured from the Carrier's online tracking, "
        f"{case_data.get('number_of_containers', '#NO.')} x empty {case_data.get('container_type', '#CONTAINER TYPES')} "
        f"container(s), No. {case_data.get('container_number', '#CONTAINER NOS')}, "
        f"gated out from the terminal at the port of {case_data.get('origin_port', '#PORT NAME')}, "
        f"{case_data.get('origin_country', '#COUNTRY')} on {case_data.get('gate_out_date', '#DATE')}."
    )

    p = doc.add_paragraph()
    p.add_run(
        f"The container(s) was returned, fully laden, to the port of "
        f"{case_data.get('origin_port', '#PORT NAME')} on {case_data.get('container_return_date', '#DATE')}, "
        f"where it was received by the Carrier, {case_data.get('carrier_name', '#NAME OF CARRIER')} for "
        f"further shipment to {case_data.get('discharge_port', '#DISPORT')}, "
        f"{case_data.get('destination_country', '#COUNTRY')} on {case_data.get('shipment_terms', '#SHIPMENT TERMS')} terms "
        f"under cover of Bill of Lading No. {case_data.get('bl_number', '#B/L NO')} "
        f"issued at {case_data.get('bl_issue_place', '#PLACE')} on {case_data.get('bl_issue_date', '#DATE')}."
    )

    p = doc.add_paragraph()
    p.add_run(
        f"The container(s) was shipped on board the carrying vessel, "
        f"M/V \"{case_data.get('vessel_name', '#NAME OF VESSEL')}\" "
        f"VOY. {case_data.get('voyage_number', '#VOYAGE NO.')} "
        f"at {case_data.get('origin_port', '#PORT NAME')} on {case_data.get('vessel_loading_date', '#DATE')}."
    )

    # Transhipment (if applicable)
    if case_data.get('has_transhipment'):
        p = doc.add_paragraph()
        p.add_run(
            f"The vessel arrived at the transhipment port of {case_data.get('transhipment_port', '#PORT')}, "
            f"{case_data.get('transhipment_country', '#COUNTRY')} on {case_data.get('transhipment_arrival_date', '#DATE')} "
            f"where the container(s) was discharged later on the same day and then further loaded "
            f"on board the on-carrying vessel, M/V \"{case_data.get('oncarrying_vessel', '#NAME OF VESSEL')}\" "
            f"VOY. {case_data.get('oncarrying_voyage', '#VOYAGE NO.')} at "
            f"{case_data.get('transhipment_port', '#PORT NAME')} on {case_data.get('transhipment_reload_date', '#DATE')}."
        )

    p = doc.add_paragraph()
    p.add_run(
        f"The vessel arrived at the final discharge port of {case_data.get('discharge_port', '#PORT')}, "
        f"{case_data.get('destination_country', '#COUNTRY')} on {case_data.get('arrival_date', '#DATE')} "
        f"where the container(s) was subsequently discharged and moved into the CY for temporary storage pending collection."
    )

    p = doc.add_paragraph()
    p.add_run(
        f"Following completion of import formalities, the container(s) were collected from the terminal "
        f"at the port on {case_data.get('collection_date', '#DATE')} for delivery to the "
        f"Consignee's {case_data.get('delivery_type', '#nominated #premises')}, "
        f"located at {case_data.get('delivery_location', '#CITY')}, "
        f"arriving on {case_data.get('delivery_date', '#DATE')}."
    )

    p = doc.add_paragraph()
    p.add_run(
        f"It was reported that at the time of the delivery of the container(s), it was "
        f"found to be in an apparent sound condition, with original shipping seal(s) still intact, "
        f"however, upon opening of the doors of container No. {case_data.get('container_number', '#CNTR NO.')}, "
        f"the receiving personnel found that {case_data.get('damage_discovery', '{REPORT IN DETAIL WHAT WAS FOUND BY THE CONSIGNEE AT THIS TIME, INCLUDE ANY PHOTOGRAPHS TAKEN BY THE CONSIGNEE AND WHAT THE CONSIGNEE DID NEXT}')}."
    )

    p = doc.add_paragraph()
    p.add_run(
        "Following discovery, the Consignee report the matter to concerned parties, as a result of which, "
        "we were requested to attend survey in order to establish nature, extent and cause of any resulting loss."
    )

    doc.add_heading('1.2    Arrangements for Survey', level=2)

    p = doc.add_paragraph()
    p.add_run(
        f"Following receipt of instructions, we immediately contacted {case_data.get('consignee_contact', '#PIC')} "
        f"of the Consignee, in order to make necessary arrangements for survey. From discussions, we understood that "
        f"{case_data.get('survey_discussion', '{REPORT IN DETAIL WHAT WAS DISCUSSED WITH THE CONSIGNEE AT THE TIME OF SURVEY ARRANGEMENTS}')}"
    )

    p = doc.add_paragraph()
    p.add_run(f"Therefore, arrangements were made to attend inspection on {case_data.get('survey_date', '#DATE')}.")

    # ============================================================
    # 2. SURVEY
    # ============================================================
    doc.add_heading('2.    SURVEY', level=1)

    doc.add_heading('2.1    Description of Goods and Packaging', level=2)

    p = doc.add_paragraph()
    p.add_run(
        f"The goods forming the subject of this claim comprised {case_data.get('number_of_packages', '#PKGS')} of "
        f"{case_data.get('goods_description', '#GOODS')} stowed in "
        f"{case_data.get('number_of_containers', '#NO.')} x empty "
        f"{case_data.get('container_type', '#CONTAINER TYPES')} container(s), "
        f"No. {case_data.get('container_number', '#CONTAINER NOS')}. "
        f"GW: {case_data.get('gross_weight', '#WEIGHT')} KGS NW: {case_data.get('net_weight', '#WEIGHT')} KGS."
    )

    p = doc.add_paragraph()
    p.add_run(case_data.get('packaging_description',
        '{REPORT IN DETAIL THE PACKAGING METHOD UTILISED, THE METHOD OF STOWAGE IN THE CONTAINER AND ANY LASHING / SECURING UTILISED - INSERT DEMONSTRATIVE PHOTOGRAPHS}'))

    # Helper to embed a group of iAuditor images with caption
    iauditor_images = case_data.get('iauditor_images', [])
    folder_id = case_data.get('id', case_data.get('case_reference', ''))

    def embed_iauditor_group(doc, category_list, caption_text=None):
        """Embed iAuditor images of given categories with a caption paragraph."""
        imgs = [img for img in iauditor_images if img.get('category') in category_list]
        if not imgs:
            return
        if caption_text:
            p = doc.add_paragraph()
            run = p.add_run(caption_text)
            run.italic = True
            run.font.size = Pt(10)
        # Sort by photo number, then group by sub_heading
        imgs = sorted(imgs, key=lambda x: (x.get('photo_number') or 9999, x.get('page', 0)))
        from itertools import groupby
        for sub, group_iter in groupby(imgs, key=lambda x: (x.get('sub_heading') or '').strip()):
            group = list(group_iter)
            if not group:
                continue
            description = (group[0].get('description') or '').strip()
            if sub:
                hp = doc.add_paragraph()
                hr = hp.add_run(sub)
                hr.bold = True
                hr.font.size = Pt(11)
            if description:
                dp = doc.add_paragraph()
                dr = dp.add_run(description)
                dr.font.size = Pt(11)

            # Render images 2 per row in a borderless table
            valid = [img for img in group if os.path.exists(
                os.path.join(UPLOAD_FOLDER, folder_id, 'iauditor_images', img['filename']))]
            for i in range(0, len(valid), 2):
                pair = valid[i:i+2]
                tbl = doc.add_table(rows=2, cols=2)
                tbl.autofit = False
                for col_idx, img in enumerate(pair):
                    img_cell = tbl.rows[0].cells[col_idx]
                    cap_cell = tbl.rows[1].cells[col_idx]
                    p_img = img_cell.paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        p_img.add_run().add_picture(
                            os.path.join(UPLOAD_FOLDER, folder_id, 'iauditor_images', img['filename']),
                            width=Inches(3.0)
                        )
                    except Exception as e:
                        print(f"Failed to embed {img['filename']}: {e}")
                    pn = img.get('photo_number')
                    cap_p = cap_cell.paragraphs[0]
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cap_p.add_run(f"Photo {pn}" if pn else "")
                    cr.italic = True
                    cr.font.size = Pt(9)
                doc.add_paragraph()

    # Embed cargo packaging photos from iAuditor.
    # If iAuditor has no dedicated packaging photos, fall back to the first
    # couple of container-interior shots (which show the stowed cargo).
    has_packaging = any(i.get('category') == 'cargo_packaging' for i in iauditor_images)
    if has_packaging:
        embed_iauditor_group(doc, ['cargo_packaging'],
            "Photographs of cargo packaging and stowage as observed during survey:")
    else:
        interior = sorted(
            [i for i in iauditor_images if i.get('category') == 'container_interior'],
            key=lambda x: (x.get('photo_number') or 9999)
        )[:2]
        if interior:
            cap = doc.add_paragraph()
            cr = cap.add_run("Photographs of cargo stowage and packaging as observed inside the container:")
            cr.italic = True
            cr.font.size = Pt(10)
            for img in interior:
                img_path = os.path.join(UPLOAD_FOLDER, folder_id, 'iauditor_images', img['filename'])
                if not os.path.exists(img_path):
                    continue
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img_path, width=Inches(4.0))
                    pn = img.get('photo_number')
                    if pn:
                        cap2 = doc.add_paragraph()
                        cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cr2 = cap2.add_run(f"Photo {pn}")
                        cr2.italic = True
                        cr2.font.size = Pt(9)
                except Exception as e:
                    print(f"Failed to embed packaging fallback {img['filename']}: {e}")

    # 2.2 Condition of Container
    doc.add_heading('2.2    Condition of Container', level=2)

    images = case_data.get('images', [])
    container_images = [img for img in images if img.get('category') == 'container']

    container_condition = case_data.get('container_condition_description', '')
    container_exterior = case_data.get('container_exterior_condition', '')
    container_interior = case_data.get('container_interior_condition', '')
    container_damages = case_data.get('container_damages_found', '')

    if container_exterior and container_exterior != 'N/A':
        p = doc.add_paragraph()
        p.add_run(f"External condition: {container_exterior}")
    if container_interior and container_interior != 'N/A':
        p = doc.add_paragraph()
        p.add_run(f"Internal condition: {container_interior}")
    if container_damages and container_damages != 'N/A':
        p = doc.add_paragraph()
        p.add_run(f"Damages found: {container_damages}")
    if container_condition and container_condition != 'N/A':
        p = doc.add_paragraph()
        p.add_run(container_condition)

    if not any(v and v != 'N/A' for v in [container_exterior, container_interior, container_damages, container_condition]):
        if container_images:
            for img in container_images:
                ai_desc = img.get('ai_description', '')
                if ai_desc:
                    p = doc.add_paragraph()
                    p.add_run(ai_desc)
        else:
            p = doc.add_paragraph()
            p.add_run("{IF AVAILABLE, FULL DESCRIPTION OF THE CONDITION OF THE CONTAINER}")

    # Embed iAuditor container exterior + interior photos
    embed_iauditor_group(doc, ['container_exterior'],
        "Container exterior photographs:")
    embed_iauditor_group(doc, ['container_interior'],
        "Container interior photographs:")

    # 2.3 Condition of Goods
    doc.add_heading('2.3    Condition of Goods', level=2)

    p = doc.add_paragraph()
    p.add_run("At the time of attendance, the goods had already been sorted and set aside by the Consignee, pending survey.")

    cargo_condition = case_data.get('cargo_condition_description', '')
    damage_details = case_data.get('damage_details', '')
    cargo_images = [img for img in images if img.get('category') == 'cargo']

    if cargo_condition and cargo_condition != 'N/A':
        p = doc.add_paragraph()
        p.add_run(cargo_condition)
    if damage_details and damage_details != 'N/A':
        p = doc.add_paragraph()
        p.add_run(damage_details)

    if not any(v and v != 'N/A' for v in [cargo_condition, damage_details]):
        if cargo_images:
            for img in cargo_images:
                ai_desc = img.get('ai_description', '')
                if ai_desc:
                    p = doc.add_paragraph()
                    p.add_run(ai_desc)
        else:
            p = doc.add_paragraph()
            p.add_run("{FULL DESCRIPTION OF THE CONDITION OF THE GOODS}")

    # Embed iAuditor cargo damage photos
    embed_iauditor_group(doc, ['cargo_condition'],
        "Photographs of cargo condition as observed during survey:")

    # 2.4 Testing
    doc.add_heading('2.4    Temperature / Chemical Testing / Moisture Testing', level=2)

    silver_nitrate = case_data.get('silver_nitrate_test', '')
    light_test = case_data.get('light_test', '')
    other_tests = case_data.get('other_tests', '')

    if silver_nitrate or light_test or other_tests:
        if silver_nitrate:
            p = doc.add_paragraph()
            p.add_run(f"Silver Nitrate Test for Chlorides: {silver_nitrate}")
        if light_test:
            p = doc.add_paragraph()
            p.add_run(f"Light Test: {light_test}")
        if other_tests:
            p = doc.add_paragraph()
            p.add_run(other_tests)
    else:
        p = doc.add_paragraph()
        p.add_run(case_data.get('testing_details',
            '{THIS SECTION SHOULD DETAIL ANY APPLICABLE TESTING CARRIED OUT}'))

    # Embed iAuditor testing photos (e.g. silver nitrate test)
    embed_iauditor_group(doc, ['testing'],
        "Photographs of testing carried out during survey:")

    # ============================================================
    # 3. DISCUSSIONS
    # ============================================================
    doc.add_heading('3.    DISCUSSIONS', level=1)

    p = doc.add_paragraph()
    p.add_run(
        f"Following survey, we discussed the Consignee's further intentions in regard to the cargo and were advised that "
        f"{case_data.get('discussions', '{THIS SECTION DETAILS DISCUSSIONS HELD DURING/FOLLOWING SURVEY ALONG WITH RECOMMENDATIONS / ACTION AGREED IN ORDER TO ESTABLISH / MINIMISE THE LOSS}.')}"
    )

    # ============================================================
    # 4. DEVELOPMENTS
    # ============================================================
    doc.add_heading('4.    DEVELOPMENTS', level=1)

    p = doc.add_paragraph()
    p.add_run(
        f"We continued to maintain contact with the Consignee and on {case_data.get('development_date', '#DATE')} "
        f"were advised that {case_data.get('developments', '{THIS SECTION DETAILS DEVELOPMENTS FOLLOWING ON FROM SURVEY (CARGO WAS SORTED / SALVAGE SALE, ETC.)}')}"
    )

    # ============================================================
    # 5. QUANTIFICATION OF LOSS
    # ============================================================
    doc.add_heading('5.    QUANTIFICATION OF LOSS', level=1)

    doc.add_heading('5.1    Loss', level=2)

    p = doc.add_paragraph()
    p.add_run(
        f"According to Commercial Invoice No. {case_data.get('invoice_number', '#NO.')} "
        f"dated {case_data.get('invoice_date', '#DATE')}, the value of the goods forming the subject "
        f"of this claim amounts to {case_data.get('claim_currency', '#CURRENCY')} "
        f"{case_data.get('claim_amount', '#AMOUNT')} {case_data.get('incoterms', '#INCOTERMS')}."
    )

    p = doc.add_paragraph()
    p.add_run(case_data.get('loss_details',
        '{FULL DETAILS OF THE LOSS UNDER REVIEW RELATING SOLELY TO THE CARGO}'))

    doc.add_heading('5.2    Additional Costs', level=2)

    p = doc.add_paragraph()
    p.add_run(case_data.get('additional_costs',
        '{DETAILS OF ANY ADDITIONAL COSTS CLAIMED (RE-PACKING / SORTING, ETC)}'))

    # ============================================================
    # 6. CAUSE OF LOSS
    # ============================================================
    doc.add_heading('6.    CAUSE OF LOSS', level=1)

    p = doc.add_paragraph()
    p.add_run(
        f"From findings during survey, we attribute the loss in this instance to "
        f"{case_data.get('cause_of_loss', '#CAUSE')}."
    )

    p = doc.add_paragraph()
    p.add_run(case_data.get('cause_explanation',
        '{PROVIDE EXPLANATION AS TO CONCLUSION AS TO THE CAUSE OF THE LOSS, AS WELL AS LIABLE PARTY TO ASSIST WITH POSSIBLE RECOVERY ACTION}'))

    # ============================================================
    # 7. PHOTOGRAPHS
    # ============================================================
    doc.add_heading('7.    PHOTOGRAPHS', level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Photographs taken at the time of survey, along with those supplied by parties concerned are embedded below."
    )

    if images:
        category_titles = {
            'container': 'Container Photographs',
            'cargo': 'Cargo / Goods Photographs',
            'vessel': 'Vessel Photographs'
        }

        photo_num = 1
        folder_name = case_data.get('id', case_data.get('case_reference', ''))

        for cat_key in ['container', 'cargo', 'vessel']:
            cat_images = [img for img in images if img.get('category') == cat_key]
            if not cat_images:
                continue

            doc.add_paragraph()
            cat_heading = doc.add_paragraph()
            run = cat_heading.add_run(category_titles[cat_key])
            run.bold = True
            run.font.size = Pt(12)

            for img in cat_images:
                image_path = os.path.join(UPLOAD_FOLDER, folder_name, 'images', img['filename'])

                if os.path.exists(image_path):
                    try:
                        img_p = doc.add_paragraph()
                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_p.add_run()
                        run.add_picture(image_path, width=Inches(4.5))
                    except Exception:
                        p = doc.add_paragraph()
                        p.add_run(f"[Image could not be embedded: {img['original_name']}]")

                # Caption
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(f"Photograph {photo_num}: {img.get('original_name', 'N/A')}")
                run.bold = True
                run.font.size = Pt(10)

                # AI Description under photo
                ai_desc = img.get('ai_description', '')
                if ai_desc:
                    desc_p = doc.add_paragraph()
                    desc_p.add_run(ai_desc)
                    desc_p.paragraph_format.left_indent = Inches(0.3)

                doc.add_paragraph()
                photo_num += 1

    doc.add_paragraph()

    # ============================================================
    # CLOSING / DISCLAIMER
    # ============================================================
    p = doc.add_paragraph()
    p.add_run(
        "This Certificate of Survey is issued, without prejudice, and subject to the terms and conditions "
        "of the relative Policy of Insurance."
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature
    p = doc.add_paragraph()
    p.add_run("for").italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    attending_surveyor_name = (
        case_data.get('attending_surveyor')
        or case_data.get('surveyor_in_charge')
        or ''
    ).strip()
    if attending_surveyor_name:
        p = doc.add_paragraph()
        r = p.add_run(attending_surveyor_name.upper())
        r.bold = True

    p = doc.add_paragraph()
    p.add_run("SURVEYOR").bold = True
    doc.add_paragraph("__________")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Copy documents enclosed:").bold = True

    enclosures = case_data.get('iauditor_enclosures', []) or []
    if enclosures:
        # Build a public URL to the iAuditor source PDF for this case
        try:
            pdf_url = url_for('serve_iauditor_pdf', case_ref=case_ref, _external=True)
        except Exception:
            pdf_url = f"/uploads/{case_ref}/iauditor_pdf"
        for name in enclosures:
            li = doc.add_paragraph(style='List Number')
            _add_hyperlink(li, pdf_url, name, size=11)

    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = footer.add_run("RESTRICTED (AUTHORISED PERSONNEL ONLY)")
    run.font.size = Pt(8)
    run.bold = True

    # Save
    filename = f"Report_{case_ref}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(REPORTS_FOLDER, filename)
    doc.save(filepath)

    return filename


# ===========================
# ROUTES
# ===========================

@app.route('/')
def index():
    if 'username' in session:
        return redirect(url_for('upload'))
    return redirect(url_for('login'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        if username in USERS and USERS[username] == password:
            session['username'] = username
            flash('Login successful!', 'success')
            return redirect(url_for('upload'))
        else:
            flash('Invalid username or password', 'error')
    return render_template('login.html')


@app.route('/logout')
def logout():
    session.pop('username', None)
    flash('You have been logged out', 'info')
    return redirect(url_for('login'))


@app.route('/email/fetch-latest-document', methods=['POST'])
def fetch_latest_email_document():
    """On-demand pull of the most recent unread M365 inbox attachment.

    Saves the file into UPLOAD_FOLDER/email_inbox/<uuid>_<name> (same root
    used by manual uploads) and returns metadata + a download URL the
    frontend uses to inject the file into the existing upload form so it
    flows through the same processing pipeline.
    """
    if 'username' not in session:
        return jsonify({'error': 'unauthorized'}), 401

    try:
        service = GraphEmailService()
        fetched = service.fetch_latest_supported_attachment()
    except EmailServiceError as e:
        return jsonify({'error': str(e)}), 500
    except Exception as e:
        return jsonify({'error': f'Graph error: {e}'}), 500

    if fetched is None:
        return jsonify({'error': 'No unread email with a supported attachment found.'}), 404

    inbox_folder = os.path.join(UPLOAD_FOLDER, 'email_inbox')
    os.makedirs(inbox_folder, exist_ok=True)
    safe_name = secure_filename(fetched.file_name)
    stored_name = f"{uuid.uuid4().hex}_{safe_name}"
    stored_path = os.path.join(inbox_folder, stored_name)
    with open(stored_path, 'wb') as f:
        f.write(fetched.content_bytes)

    return jsonify({
        'file_name': fetched.file_name,
        'stored_name': stored_name,
        'content_type': fetched.content_type,
        'size_bytes': len(fetched.content_bytes),
        'subject': fetched.subject,
        'sender': fetched.sender,
        'received_at': fetched.received_at,
        'download_url': url_for('serve_email_inbox_file', stored_name=stored_name),
    })


@app.route('/email/inbox/<stored_name>')
def serve_email_inbox_file(stored_name):
    if 'username' not in session:
        return redirect(url_for('login'))
    safe = secure_filename(stored_name)
    return send_from_directory(os.path.join(UPLOAD_FOLDER, 'email_inbox'), safe)


@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if 'username' not in session:
        return redirect(url_for('login'))

    if request.method == 'POST':
        # Generate unique case ID
        case_id = f"CASE-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{uuid.uuid4().hex[:6]}"
        case_reference = request.form.get('case_reference', '').strip()
        if not case_reference:
            case_reference = case_id

        container_number = request.form.get('container_number', '')
        bl_number = request.form.get('bl_number', '')
        goods_description = request.form.get('goods_description', '')
        shipper = request.form.get('shipper', '')
        consignee = request.form.get('consignee', '')

        # Create case folder using unique ID
        case_folder = os.path.join(UPLOAD_FOLDER, case_id)
        os.makedirs(case_folder, exist_ok=True)

        # Handle file uploads
        documents = {}
        file_types = ['bill_of_lading', 'commercial_invoice', 'packing_list', 'iauditor_report']

        for file_type in file_types:
            if file_type in request.files:
                file = request.files[file_type]
                if file and file.filename and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(case_folder, f"{file_type}_{filename}")
                    file.save(filepath)
                    documents[file_type] = filename

        # Collect all extra extracted fields from hidden form inputs
        extra_fields = [
            'principal_reference', 'surveyor_in_charge', 'attending_surveyor',
            'survey_date', 'survey_location', 'other_parties',
            'container_type', 'number_of_containers', 'bl_issue_place', 'bl_issue_date',
            'carrier_name', 'vessel_name', 'voyage_number',
            'origin_port', 'origin_country', 'discharge_port', 'destination_country',
            'transhipment_port', 'arrival_date', 'delivery_date', 'delivery_location',
            'gate_out_date', 'container_return_date', 'vessel_loading_date',
            'shipment_terms', 'incoterms', 'has_transhipment',
            'number_of_packages', 'gross_weight', 'net_weight',
            'container_condition_description', 'container_exterior_condition',
            'container_interior_condition', 'container_damages_found',
            'packaging_description', 'cargo_condition_description',
            'quantities_offloaded', 'quantities_inspected', 'damage_details',
            'silver_nitrate_test', 'light_test', 'other_tests',
            'cause_of_loss', 'cause_explanation', 'discussions', 'damage_discovery',
            'invoice_number', 'invoice_date', 'claim_currency', 'claim_amount',
            'notice_of_loss_date', 'consignee_contact', 'survey_discussion',
            'delivery_type', 'collection_date', 'transit_routing',
        ]

        case_data = {
            'id': case_id,
            'case_reference': case_reference,
            'container_number': container_number,
            'bl_number': bl_number,
            'goods_description': goods_description,
            'shipper': shipper,
            'consignee': consignee,
            'documents': documents,
            'images': [],
            'created_by': session['username'],
            'created_at': datetime.now().isoformat(),
            'status': 'pending'
        }

        # Add all extra fields from form (populated by extraction JS)
        for field in extra_fields:
            val = request.form.get(field, '').strip()
            if val:
                # Convert has_transhipment to boolean
                if field == 'has_transhipment':
                    case_data[field] = val.lower() in ('true', 'yes', '1')
                else:
                    case_data[field] = val

        # Extract images from iAuditor PDF if uploaded
        if 'iauditor_report' in documents:
            iauditor_path = os.path.join(case_folder, f"iauditor_report_{documents['iauditor_report']}")
            if os.path.exists(iauditor_path):
                iauditor_images_folder = os.path.join(case_folder, 'iauditor_images')
                extracted_imgs = extract_iauditor_images(iauditor_path, iauditor_images_folder)
                # Convert raw checklist text into clean professional narratives
                try:
                    extracted_imgs = narrate_iauditor_sections(extracted_imgs)
                except Exception as e:
                    print(f"Narration step failed: {e}")
                case_data['iauditor_images'] = extracted_imgs
                try:
                    case_data['iauditor_enclosures'] = extract_iauditor_file_summary(iauditor_path)
                except Exception as e:
                    print(f"Enclosure extraction failed: {e}")
                    case_data['iauditor_enclosures'] = []
                print(f"Extracted {len(extracted_imgs)} images from iAuditor PDF")

        cases = load_case_data()
        cases.append(case_data)
        save_case_data(cases)

        flash(f'Case {case_reference} created successfully!', 'success')
        return redirect(url_for('images_page', case_id=case_id))

    return render_template('upload.html', username=session['username'])


@app.route('/images/<case_id>')
def images_page(case_id):
    if 'username' not in session:
        return redirect(url_for('login'))

    cases = load_case_data()
    case_data = find_case_by_id(cases, case_id)

    # Fallback for old data without 'id'
    if not case_data:
        case_data = find_case_by_ref(cases, case_id)

    if not case_data:
        flash('Case not found', 'error')
        return redirect(url_for('upload'))

    return render_template('images.html', case=case_data, username=session['username'])


@app.route('/uploads/<case_ref>/iauditor_pdf')
def serve_iauditor_pdf(case_ref):
    """Serve the original iAuditor PDF used for a case."""
    if 'username' not in session:
        return redirect(url_for('login'))
    case_dir = os.path.join(UPLOAD_FOLDER, case_ref)
    if not os.path.isdir(case_dir):
        return "Case not found", 404
    for fname in os.listdir(case_dir):
        if fname.startswith('iauditor_report_') and fname.lower().endswith('.pdf'):
            return send_from_directory(case_dir, fname)
    return "iAuditor report not found", 404


@app.route('/uploads/<case_ref>/images/<filename>')
def serve_case_image(case_ref, filename):
    """Serve uploaded images for display."""
    if 'username' not in session:
        return redirect(url_for('login'))
    # Check user-uploaded images
    image_dir = os.path.join(UPLOAD_FOLDER, case_ref, 'images')
    if os.path.exists(os.path.join(image_dir, filename)):
        return send_from_directory(image_dir, filename)
    # Check iAuditor extracted images
    iauditor_dir = os.path.join(UPLOAD_FOLDER, case_ref, 'iauditor_images')
    if os.path.exists(os.path.join(iauditor_dir, filename)):
        return send_from_directory(iauditor_dir, filename)
    return "Image not found", 404


@app.route('/api/upload-images/<case_id>', methods=['POST'])
def api_upload_images(case_id):
    if 'username' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    cases = load_case_data()
    case_data = find_case_by_id(cases, case_id)
    if not case_data:
        case_data = find_case_by_ref(cases, case_id)

    if not case_data:
        return jsonify({'error': 'Case not found'}), 404

    images = request.files.getlist('images')
    categories = request.form.getlist('categories')

    if not images or not categories:
        return jsonify({'error': 'No images or categories provided'}), 400

    # Use the case ID (unique) for folder, not case_reference
    folder_name = case_data.get('id', case_data.get('case_reference', case_id))
    images_folder = os.path.join(UPLOAD_FOLDER, folder_name, 'images')
    os.makedirs(images_folder, exist_ok=True)

    if 'images' not in case_data:
        case_data['images'] = []

    results = []
    for i, (image_file, category) in enumerate(zip(images, categories)):
        if not image_file or not image_file.filename:
            continue
        if not allowed_image(image_file.filename):
            results.append({'filename': image_file.filename, 'error': 'Invalid file type'})
            continue

        original_name = secure_filename(image_file.filename)
        timestamp = datetime.now().strftime('%H%M%S')
        filename = f"{category}_{timestamp}_{i}_{original_name}"
        filepath = os.path.join(images_folder, filename)
        image_file.save(filepath)

        analysis = analyze_damage_image(filepath, category)

        image_record = {
            'filename': filename,
            'original_name': original_name,
            'category': category,
            'ai_description': analysis.get('description', '') if analysis.get('success') else '',
            'ai_error': analysis.get('error', '') if not analysis.get('success') else '',
            'uploaded_at': datetime.now().isoformat()
        }

        case_data['images'].append(image_record)
        results.append({
            'filename': filename,
            'category': category,
            'analyzed': analysis.get('success', False),
        })

    save_case_data(cases)

    analyzed_count = sum(1 for r in results if r.get('analyzed'))
    return jsonify({
        'success': True,
        'message': f'{len(results)} image(s) uploaded, {analyzed_count} analyzed by AI.',
        'results': results
    })


@app.route('/generate/<case_id>')
def generate_report_page(case_id):
    if 'username' not in session:
        return redirect(url_for('login'))

    cases = load_case_data()
    case_data = find_case_by_id(cases, case_id)
    if not case_data:
        case_data = find_case_by_ref(cases, case_id)

    if not case_data:
        flash('Case not found', 'error')
        return redirect(url_for('upload'))

    return render_template('generate.html', case=case_data, username=session['username'])


@app.route('/api/generate-report/<case_id>', methods=['POST'])
def api_generate_report(case_id):
    if 'username' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    cases = load_case_data()
    case_data = find_case_by_id(cases, case_id)
    if not case_data:
        case_data = find_case_by_ref(cases, case_id)

    if not case_data:
        return jsonify({'error': 'Case not found'}), 404

    try:
        case_data['report_date'] = datetime.now().strftime('%d %B %Y')
        filename = generate_full_report(case_data)

        case_data['status'] = 'completed'
        case_data['report_file'] = filename
        case_data['completed_at'] = datetime.now().isoformat()
        save_case_data(cases)

        return jsonify({
            'success': True,
            'message': 'Report generated successfully',
            'filename': filename,
            'download_url': url_for('download_report', filename=filename)
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/download/<filename>')
def download_report(filename):
    if 'username' not in session:
        return redirect(url_for('login'))

    filepath = os.path.join(REPORTS_FOLDER, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    else:
        flash('Report not found', 'error')
        return redirect(url_for('dashboard'))


@app.route('/dashboard')
def dashboard():
    if 'username' not in session:
        return redirect(url_for('login'))

    cases = load_case_data()
    cases.sort(key=lambda x: x.get('created_at', ''), reverse=True)

    stats = {
        'total': len(cases),
        'pending': len([c for c in cases if c.get('status') == 'pending']),
        'completed': len([c for c in cases if c.get('status') == 'completed'])
    }

    return render_template('dashboard.html', cases=cases, stats=stats, username=session['username'])


@app.route('/api/extract-data', methods=['POST'])
def api_extract_data():
    if 'username' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        documents = {}
        file_types = ['bill_of_lading', 'commercial_invoice', 'packing_list', 'iauditor_report']
        temp_folder = os.path.join(UPLOAD_FOLDER, 'temp_extraction')
        os.makedirs(temp_folder, exist_ok=True)

        for file_type in file_types:
            if file_type in request.files:
                file = request.files[file_type]
                if file and file.filename and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(temp_folder, f"{file_type}_{filename}")
                    file.save(filepath)
                    documents[file_type] = filename

        if not documents:
            return jsonify({'error': 'No valid documents uploaded', 'extracted': False}), 400

        extracted_data = process_uploaded_documents(temp_folder, documents)

        import shutil
        if os.path.exists(temp_folder):
            shutil.rmtree(temp_folder)

        return jsonify(extracted_data)

    except Exception as e:
        return jsonify({'error': f'Extraction failed: {str(e)}', 'extracted': False}), 500


@app.route('/contact-admin')
def contact_admin():
    return render_template('contact.html')


if __name__ == '__main__':
    print("\n" + "=" * 60)
    print("CONTAINER INSPECTION REPORT SYSTEM")
    print("=" * 60)
    print("\nDefault Login Credentials:")
    print("  Username: admin    Password: admin123")
    print("  Username: surveyor Password: survey123")
    print("  Username: demo     Password: demo123")
    print("\nServer starting...")
    print("Access the application at: http://localhost:5001")
    print("=" * 60 + "\n")

    app.run(debug=True, host='0.0.0.0', port=5001)
